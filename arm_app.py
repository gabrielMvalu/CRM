
import streamlit as st
from docx import Document
import pandas as pd
import numpy as np
import pydeck as pdk


   chart_data = pd.DataFrame(np.random.randn(30, 3), columns=["a", "b", "c"])
   
   st.vega_lite_chart(
      chart_data,
      {
        "$schema": "https://vega.github.io/schema/vega-lite/v5.json",
        "config": {"view": {"stroke": ""}},
        "width": 800,
        "height": 200,
        "data": {
          "values": [
            {"country": "Great Britain", "animal": "pigs"},
            {"country": "Great Britain", "animal": "pigs"},
            {"country": "Great Britain", "animal": "cattle"},
            {"country": "Great Britain", "animal": "cattle"},
            {"country": "Great Britain", "animal": "cattle"},
            {"country": "Great Britain", "animal": "sheep"},
            {"country": "Great Britain", "animal": "sheep"},
            {"country": "Great Britain", "animal": "sheep"},
            {"country": "Great Britain", "animal": "sheep"},
            {"country": "Great Britain", "animal": "sheep"},
            {"country": "Great Britain", "animal": "sheep"},
            {"country": "Great Britain", "animal": "sheep"},
            {"country": "Great Britain", "animal": "sheep"},
            {"country": "Great Britain", "animal": "sheep"},
            {"country": "Great Britain", "animal": "sheep"},
            {"country": "United States", "animal": "pigs"},
            {"country": "United States", "animal": "pigs"},
            {"country": "United States", "animal": "pigs"},
            {"country": "United States", "animal": "pigs"},
            {"country": "United States", "animal": "pigs"},
            {"country": "United States", "animal": "pigs"},
            {"country": "United States", "animal": "cattle"},
            {"country": "United States", "animal": "cattle"},
            {"country": "United States", "animal": "cattle"},
            {"country": "United States", "animal": "cattle"},
            {"country": "United States", "animal": "cattle"},
            {"country": "United States", "animal": "cattle"},
            {"country": "United States", "animal": "cattle"},
            {"country": "United States", "animal": "cattle"},
            {"country": "United States", "animal": "cattle"},
            {"country": "United States", "animal": "sheep"},
            {"country": "United States", "animal": "sheep"},
            {"country": "United States", "animal": "sheep"},
            {"country": "United States", "animal": "sheep"},
            {"country": "United States", "animal": "sheep"},
            {"country": "United States", "animal": "sheep"},
            {"country": "United States", "animal": "sheep"}
          ]
        },
        "transform": [
          {
            "calculate": "{'cattle': '🐄', 'pigs': '🐖', 'sheep': '🐏'}[datum.animal]",
            "as": "emoji"
          },
          {"window": [{"op": "rank", "as": "rank"}], "groupby": ["country", "animal"]}
        ],
        "mark": {"type": "text", "baseline": "middle"},
        "encoding": {
          "x": {"field": "rank", "type": "ordinal"},
          "y": {"field": "animal", "type": "nominal"},
          "row": {"field": "country", "header": {"title": ""}},
          "text": {"field": "emoji", "type": "nominal"},
          "size": {"value": 65}
        }
      }

   )


   chart_data = pd.DataFrame(
      np.random.randn(10, 2) / [50, 50] + [44.02, 23.34],
      columns=['lat', 'lon'])
   
   st.pydeck_chart(pdk.Deck(
       map_style=None,
       initial_view_state=pdk.ViewState(
           latitude=44.02,
           longitude=23.34,
           zoom=11,
           pitch=50,
       ),
       layers=[
           pdk.Layer(
              'HexagonLayer',
              data=chart_data,
              get_position='[lon, lat]',
              radius=200,
              elevation_scale=4,
              elevation_range=[0, 1000],
              pickable=True,
              extruded=True,
           ),
           pdk.Layer(
               'ScatterplotLayer',
               data=chart_data,
               get_position='[lon, lat]',
               get_color='[200, 30, 0, 160]',
               get_radius=200,
           ),
       ],
   ))




def read_docx(file):
    doc = Document(file)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip() != '']
    return paragraphs

st.sidebar.title("Încărcare Document")
uploaded_file = st.sidebar.file_uploader("Alege un fișier .docx", type="docx")

if uploaded_file is not None:
    # Încarcă și prelucrează documentul o singură dată
    if 'paragraphs' not in st.session_state or st.session_state.uploaded_file != uploaded_file:
        st.session_state.paragraphs = read_docx(uploaded_file)
        st.session_state.paragraph_states = ['🔴' for _ in st.session_state.paragraphs]  # Inițializează toți indicatorii ca roșii
        st.session_state.uploaded_file = uploaded_file

    for i, original_text in enumerate(st.session_state.paragraphs):
        col1, col2, col3 = st.columns([1, 20, 5])
        with col1:
            st.markdown(st.session_state.paragraph_states[i], unsafe_allow_html=True)
        with col2:
            # Folosește un key unic pentru fiecare text_area pentru a putea detecta modificările
            user_input = st.text_area(f"Paragraful {i+1}", value=original_text, key=f'para_{i}')
        with col3:
            save_button = st.button('Salvează', key=f'save_{i}')
            if save_button:
                if user_input != original_text:
                    st.session_state.paragraph_states[i] = '🔵'  # Modificat și diferit de original
                else:
                    st.session_state.paragraph_states[i] = '🟢'  # Modificat dar identic cu originalul
                st.experimental_rerun()

else:
    st.write("Vă rugăm să încărcați un document .docx în meniul din stânga.")
